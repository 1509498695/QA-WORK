from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt
from PIL import Image, ImageDraw, ImageFont
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.cidfonts import UnicodeCIDFont
from reportlab.pdfgen import canvas


ROOT = Path(__file__).resolve().parent
SOURCES = ROOT / "fixtures" / "sources"


def fixture_font(size: int) -> ImageFont.FreeTypeFont | ImageFont.ImageFont:
    for candidate in (
        Path(r"C:\Windows\Fonts\msyh.ttc"),
        Path(r"C:\Windows\Fonts\simhei.ttf"),
        Path(r"C:\Windows\Fonts\arial.ttf"),
    ):
        if candidate.exists():
            return ImageFont.truetype(str(candidate), size=size)
    return ImageFont.load_default()


def create_image() -> Path:
    path = SOURCES / "reward-state.png"
    image = Image.new("RGB", (760, 300), "#F4F1E8")
    draw = ImageDraw.Draw(image)
    title_font = fixture_font(34)
    body_font = fixture_font(25)
    draw.rounded_rectangle((24, 24, 736, 276), radius=24, fill="#DCE8E1", outline="#2F5D50", width=4)
    draw.text((60, 58), "周年庆宝箱", font=title_font, fill="#24483F")
    draw.text((60, 125), "首次开启：100 金币", font=body_font, fill="#20322C")
    draw.text((60, 178), "背包已满：奖励发送至邮件", font=body_font, fill="#20322C")
    image.save(path, format="PNG")
    return path


def create_docx(image_path: Path) -> Path:
    path = SOURCES / "anniversary-chest.docx"
    document = Document()
    section = document.sections[0]
    section.header.paragraphs[0].text = "SAMO 周年庆策划案"
    section.footer.paragraphs[0].text = "本地测试夹具"
    title = document.add_heading("周年庆宝箱需求", level=1)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    document.add_paragraph("活动开启且玩家达到 10 级后，可在周年庆入口开启宝箱。")
    document.add_paragraph("首次开启获得 100 金币；当天重复开启不再发奖；每日 00:00 刷新。")
    document.add_paragraph("参考链接仅登记不访问：https://example.invalid/design-detail")
    table = document.add_table(rows=1, cols=3)
    table.style = "Table Grid"
    headers = table.rows[0].cells
    headers[0].text = "状态"
    headers[1].text = "操作"
    headers[2].text = "预期"
    for values in (
        ("首次", "开启宝箱", "获得 100 金币"),
        ("重复", "再次开启", "不重复发奖"),
        ("背包已满", "开启宝箱", "奖励进入邮件"),
    ):
        cells = table.add_row().cells
        for cell, value in zip(cells, values, strict=True):
            cell.text = value
    document.add_picture(str(image_path), width=Inches(5.8))
    caption = document.add_paragraph("图 1：宝箱奖励状态示意")
    caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for paragraph in document.paragraphs:
        for run in paragraph.runs:
            run.font.name = "Microsoft YaHei"
            run.font.size = Pt(10.5)
    document.save(path)
    return path


def create_pdf() -> Path:
    path = SOURCES / "anniversary-rules.pdf"
    pdfmetrics.registerFont(UnicodeCIDFont("STSong-Light"))
    sheet = canvas.Canvas(str(path), pagesize=(595, 842))
    sheet.setTitle("周年庆规则补充")
    sheet.setFont("STSong-Light", 20)
    sheet.drawString(64, 770, "周年庆规则补充")
    sheet.setFont("STSong-Light", 12)
    lines = [
        "1. 活动每日 00:00 刷新宝箱首次奖励状态。",
        "2. 背包已满时，奖励通过邮件发放。",
        "3. 活动结束后入口关闭，未领取邮件仍可领取。",
        "4. 外部链接仅登记：https://example.invalid/pdf-link",
    ]
    y = 720
    for line in lines:
        sheet.drawString(72, y, line)
        y -= 34
    sheet.setStrokeColorRGB(47 / 255, 93 / 255, 80 / 255)
    sheet.rect(64, 500, 460, 100, stroke=1, fill=0)
    sheet.drawString(82, 565, "边界：等级 9 不可见；等级 10 可见。")
    sheet.drawString(82, 530, "异常：断线重连后保留当天已开启状态。")
    sheet.showPage()
    sheet.save()
    return path


def create_text_sources() -> tuple[Path, Path]:
    markdown = SOURCES / "anniversary-notes.md"
    markdown.write_text(
        "# 周年庆宝箱补充\n\n- 主入口：活动中心。\n- 关闭时间后不可再次开启。\n- 外链仅登记：https://example.invalid/md\n",
        encoding="utf-8",
    )
    text = SOURCES / "anniversary-copy.txt"
    text.write_text("入口文案：周年庆宝箱\n按钮文案：开启\n错误提示：今日奖励已领取\n", encoding="utf-8")
    return markdown, text


def main() -> None:
    SOURCES.mkdir(parents=True, exist_ok=True)
    image_path = create_image()
    create_docx(image_path)
    create_pdf()
    create_text_sources()


if __name__ == "__main__":
    main()
