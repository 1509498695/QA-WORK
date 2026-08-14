from __future__ import annotations

import argparse
import hashlib
import json
import os
import re
import shutil
import subprocess
import sys
import tempfile
import zipfile
from datetime import datetime, timezone
from pathlib import Path
from typing import Any, Iterable


SUPPORTED_EXTENSIONS = {".docx", ".pdf", ".xlsx", ".md", ".txt", ".png", ".jpg", ".jpeg"}
IMAGE_EXTENSIONS = {".png", ".jpg", ".jpeg"}
URL_RE = re.compile(r"https?://[^\s<>\]\[\)）}]+", re.IGNORECASE)
REVIEW_SHEET_RE = re.compile(r"cp|反馈", re.IGNORECASE)


class SourcePacketError(ValueError):
    pass


def canonical_json(value: Any) -> str:
    return json.dumps(value, ensure_ascii=False, sort_keys=True, separators=(",", ":"))


def sha256_bytes(value: bytes) -> str:
    return hashlib.sha256(value).hexdigest()


def sha256_file(path: Path) -> str:
    digest = hashlib.sha256()
    with path.open("rb") as stream:
        for block in iter(lambda: stream.read(1024 * 1024), b""):
            digest.update(block)
    return digest.hexdigest()


def write_json(path: Path, payload: dict[str, Any]) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)
    content = json.dumps(payload, ensure_ascii=False, indent=2) + "\n"
    with tempfile.NamedTemporaryFile(
        "w",
        encoding="utf-8",
        dir=path.parent,
        prefix=f".{path.name}.",
        suffix=".tmp",
        delete=False,
    ) as stream:
        stream.write(content)
        temporary = Path(stream.name)
    temporary.replace(path)


def clean_text(value: Any) -> str:
    text = str(value if value is not None else "")
    text = text.replace("\r\n", "\n").replace("\r", "\n")
    return "\n".join(line.rstrip() for line in text.split("\n")).strip()


def unique_links(values: Iterable[str]) -> list[str]:
    links: set[str] = set()
    for value in values:
        links.update(URL_RE.findall(str(value or "")))
    return sorted(links)


def content_unit(source_ref: str, kind: str, text: str, **extra: Any) -> dict[str, Any]:
    payload: dict[str, Any] = {"source_ref": source_ref, "kind": kind, "text": clean_text(text)}
    payload.update(extra)
    return payload


def evidence_record(
    evidence_id: str,
    source_id: str,
    source_ref: str,
    kind: str,
    relative_path: str,
    sha256: str,
    **extra: Any,
) -> dict[str, Any]:
    payload: dict[str, Any] = {
        "evidence_id": evidence_id,
        "source_id": source_id,
        "source_ref": source_ref,
        "kind": kind,
        "path": relative_path.replace("\\", "/"),
        "sha256": sha256,
        "visual_review_required": True,
        "review_status": "pending",
        "observations": [],
    }
    payload.update(extra)
    return payload


def save_evidence_bytes(
    data: bytes,
    evidence_dir: Path,
    source_id: str,
    kind: str,
    position: int,
    suffix: str,
    source_ref: str,
    **extra: Any,
) -> dict[str, Any]:
    safe_suffix = suffix.lower() if suffix.lower() in {".png", ".jpg", ".jpeg", ".gif", ".bmp", ".tif", ".tiff", ".webp"} else ".bin"
    name = f"{source_id.lower()}-{kind}-{position:03d}{safe_suffix}"
    path = evidence_dir / name
    path.write_bytes(data)
    return evidence_record(
        evidence_id=f"EVID-{source_id[4:]}-{position:04d}",
        source_id=source_id,
        source_ref=source_ref,
        kind=kind,
        relative_path=f"evidence/{name}",
        sha256=sha256_bytes(data),
        **extra,
    )


def base_file_record(source_id: str, path: Path) -> dict[str, Any]:
    return {
        "source_id": source_id,
        "path": str(path.resolve()),
        "file_name": path.name,
        "extension": path.suffix.lower(),
        "size_bytes": path.stat().st_size,
        "sha256": sha256_file(path),
        "status": "readable",
        "content_units": [],
        "evidence": [],
        "excluded_items": [],
        "external_links": [],
        "warnings": [],
    }


def decode_text_file(path: Path) -> tuple[str, str]:
    raw = path.read_bytes()
    for encoding in ("utf-8-sig", "utf-8", "gb18030", "utf-16"):
        try:
            return raw.decode(encoding), encoding
        except UnicodeDecodeError:
            continue
    raise SourcePacketError("文本编码无法识别")


def extract_text(path: Path, record: dict[str, Any]) -> None:
    text, encoding = decode_text_file(path)
    record["encoding"] = encoding
    blocks: list[tuple[int, str]] = []
    start = 1
    buffer: list[str] = []
    for number, line in enumerate(text.replace("\r\n", "\n").replace("\r", "\n").split("\n"), 1):
        if line.strip():
            if not buffer:
                start = number
            buffer.append(line)
        elif buffer:
            blocks.append((start, "\n".join(buffer)))
            buffer = []
    if buffer:
        blocks.append((start, "\n".join(buffer)))
    for start_line, block in blocks:
        record["content_units"].append(
            content_unit(f"{record['source_id']}#line-{start_line}", "text_block", block, start_line=start_line)
        )
    record["external_links"] = unique_links([text])
    if not clean_text(text):
        record["status"] = "partial"
        record["warnings"].append("文件没有可读文字")


def iter_docx_table_rows(table: Any) -> Iterable[str]:
    for row in table.rows:
        yield " | ".join(clean_text(cell.text) for cell in row.cells)


def extract_docx(path: Path, record: dict[str, Any], evidence_dir: Path) -> None:
    try:
        from docx import Document
    except ImportError as exc:
        raise SourcePacketError("缺少 python-docx") from exc

    document = Document(str(path))
    for index, paragraph in enumerate(document.paragraphs, 1):
        text = clean_text(paragraph.text)
        if text:
            record["content_units"].append(
                content_unit(f"{record['source_id']}#paragraph-{index}", "paragraph", text)
            )
    for table_index, table in enumerate(document.tables, 1):
        for row_index, text in enumerate(iter_docx_table_rows(table), 1):
            if clean_text(text.replace("|", "")):
                record["content_units"].append(
                    content_unit(
                        f"{record['source_id']}#table-{table_index}-row-{row_index}",
                        "table_row",
                        text,
                        table_index=table_index,
                        row_index=row_index,
                    )
                )
    header_footer_index = 0
    for section_index, section in enumerate(document.sections, 1):
        for area_name, area in (("header", section.header), ("footer", section.footer)):
            for paragraph_index, paragraph in enumerate(area.paragraphs, 1):
                text = clean_text(paragraph.text)
                if text:
                    header_footer_index += 1
                    record["content_units"].append(
                        content_unit(
                            f"{record['source_id']}#{area_name}-{section_index}-{paragraph_index}",
                            area_name,
                            text,
                        )
                    )
            for table_index, table in enumerate(area.tables, 1):
                for row_index, text in enumerate(iter_docx_table_rows(table), 1):
                    if clean_text(text.replace("|", "")):
                        header_footer_index += 1
                        record["content_units"].append(
                            content_unit(
                                f"{record['source_id']}#{area_name}-table-{section_index}-{table_index}-row-{row_index}",
                                f"{area_name}_table_row",
                                text,
                            )
                        )

    native_links: list[str] = []
    for relation in document.part.rels.values():
        if getattr(relation, "is_external", False):
            target = str(getattr(relation, "target_ref", ""))
            if target.startswith(("http://", "https://")):
                native_links.append(target)

    with zipfile.ZipFile(path) as archive:
        media = sorted(name for name in archive.namelist() if name.startswith("word/media/") and not name.endswith("/"))
        for position, member in enumerate(media, 1):
            source_ref = f"{record['source_id']}#embedded-image-{position}"
            item = save_evidence_bytes(
                archive.read(member),
                evidence_dir,
                record["source_id"],
                "docx-image",
                position,
                Path(member).suffix,
                source_ref,
                container_member=member,
            )
            record["evidence"].append(item)
    all_text = [unit["text"] for unit in record["content_units"]]
    record["external_links"] = sorted(set(native_links) | set(unique_links(all_text)))
    record["visual_review_required"] = True
    record["warnings"].append("DOCX 必须逐页渲染复核文本框、浮动对象和排版关系")
    if not record["content_units"] and not record["evidence"]:
        record["status"] = "partial"
        record["warnings"].append("DOCX 未提取到文字或图片")


def command_for_batch(path: Path, arguments: list[str]) -> list[str]:
    if path.suffix.lower() not in {".cmd", ".bat"}:
        return [str(path), *arguments]
    command = subprocess.list2cmdline([str(path), *arguments])
    return [os.environ.get("COMSPEC", "cmd.exe"), "/d", "/s", "/c", command]


def render_pdf(path: Path, pdftoppm: Path, prefix: Path) -> list[Path]:
    result = subprocess.run(
        command_for_batch(pdftoppm, ["-png", "-r", "150", str(path), str(prefix)]),
        check=False,
        capture_output=True,
        text=True,
        timeout=180,
    )
    if result.returncode != 0:
        message = clean_text(result.stderr or result.stdout) or f"退出码 {result.returncode}"
        raise SourcePacketError(f"PDF 渲染失败：{message}")
    return sorted(prefix.parent.glob(prefix.name + "-*.png"))


def extract_pdf(path: Path, record: dict[str, Any], evidence_dir: Path, pdftoppm: Path | None) -> None:
    try:
        from pypdf import PdfReader
    except ImportError as exc:
        raise SourcePacketError("缺少 pypdf") from exc

    reader = PdfReader(str(path))
    native_links: list[str] = []
    for page_index, page in enumerate(reader.pages, 1):
        text = clean_text(page.extract_text() or "")
        if text:
            record["content_units"].append(
                content_unit(f"{record['source_id']}#page-{page_index}", "pdf_page_text", text, page=page_index)
            )
        annotations = page.get("/Annots") or []
        for annotation_ref in annotations:
            try:
                annotation = annotation_ref.get_object()
                action = annotation.get("/A") or {}
                uri = action.get("/URI")
                if uri:
                    native_links.append(str(uri))
            except Exception:
                record["warnings"].append(f"第 {page_index} 页存在无法解析的批注")
    record["page_count"] = len(reader.pages)
    all_text = [unit["text"] for unit in record["content_units"]]
    record["external_links"] = sorted(set(native_links) | set(unique_links(all_text)))
    record["visual_review_required"] = True
    if pdftoppm is None:
        record["status"] = "partial"
        record["warnings"].append("未提供 pdftoppm，PDF 页面尚未渲染")
        return
    if not pdftoppm.exists():
        record["status"] = "partial"
        record["warnings"].append(f"pdftoppm 不存在：{pdftoppm}")
        return
    with tempfile.TemporaryDirectory(prefix="qa-case-xlsx-pdf-") as temporary:
        pages = render_pdf(path, pdftoppm, Path(temporary) / "page")
        if len(pages) != len(reader.pages):
            record["status"] = "partial"
            record["warnings"].append(
                f"PDF 渲染页数不一致：expected={len(reader.pages)} actual={len(pages)}"
            )
        for page_index, page_path in enumerate(pages, 1):
            source_ref = f"{record['source_id']}#page-{page_index}"
            item = save_evidence_bytes(
                page_path.read_bytes(),
                evidence_dir,
                record["source_id"],
                "pdf-page",
                page_index,
                ".png",
                source_ref,
                page=page_index,
            )
            record["evidence"].append(item)
    if not record["content_units"]:
        record["status"] = "partial"
        record["warnings"].append("PDF 未提取到文字，必须通过页面视觉读取")


def extract_xlsx_media(path: Path, record: dict[str, Any], evidence_dir: Path) -> None:
    with zipfile.ZipFile(path) as archive:
        media = sorted(name for name in archive.namelist() if name.startswith("xl/media/") and not name.endswith("/"))
        for position, member in enumerate(media, 1):
            source_ref = f"{record['source_id']}#embedded-image-{position}"
            item = save_evidence_bytes(
                archive.read(member),
                evidence_dir,
                record["source_id"],
                "xlsx-image",
                position,
                Path(member).suffix,
                source_ref,
                container_member=member,
            )
            record["evidence"].append(item)


def extract_xlsx(
    path: Path,
    record: dict[str, Any],
    evidence_dir: Path,
    include_hidden_sheets: bool,
    include_review_sheets: bool,
) -> None:
    try:
        from openpyxl import load_workbook
    except ImportError as exc:
        raise SourcePacketError("缺少 openpyxl，无法读取 XLSX") from exc

    workbook = load_workbook(filename=path, read_only=False, data_only=False, keep_links=False)
    try:
        for sheet in workbook.worksheets:
            reasons: list[str] = []
            if sheet.sheet_state != "visible" and not include_hidden_sheets:
                reasons.append(f"Sheet 状态为 {sheet.sheet_state}")
            if REVIEW_SHEET_RE.search(sheet.title) and not include_review_sheets:
                reasons.append("Sheet 名称命中 CP/反馈默认排除规则")
            if reasons:
                record["excluded_items"].append(
                    {"kind": "sheet", "name": sheet.title, "reason": "；".join(reasons)}
                )
                continue
            for row in sheet.iter_rows():
                for cell in row:
                    if cell.value is None:
                        continue
                    text = clean_text(cell.value)
                    if not text:
                        continue
                    ref = f"{record['source_id']}#sheet-{sheet.title}!{cell.coordinate}"
                    record["content_units"].append(
                        content_unit(
                            ref,
                            "xlsx_cell",
                            text,
                            sheet=sheet.title,
                            cell=cell.coordinate,
                            data_type=cell.data_type,
                        )
                    )
                    hyperlink = getattr(cell, "hyperlink", None)
                    target = getattr(hyperlink, "target", None) if hyperlink else None
                    if target and str(target).startswith(("http://", "https://")):
                        record["external_links"].append(str(target))
        record["external_links"] = sorted(
            set(record["external_links"]) | set(unique_links(unit["text"] for unit in record["content_units"]))
        )
    finally:
        workbook.close()
    extract_xlsx_media(path, record, evidence_dir)
    if record["evidence"]:
        record["visual_review_required"] = True
        record["warnings"].append("XLSX 嵌入图片必须逐张视觉复核并补充 Sheet/单元格语义定位")
    if not record["content_units"] and not record["evidence"]:
        record["status"] = "partial"
        record["warnings"].append("纳入范围内没有可读单元格或图片")


def extract_image(path: Path, record: dict[str, Any], evidence_dir: Path) -> None:
    try:
        from PIL import Image
    except ImportError as exc:
        raise SourcePacketError("缺少 Pillow") from exc

    with Image.open(path) as image:
        width, height = image.size
        image_format = image.format
    data = path.read_bytes()
    source_ref = f"{record['source_id']}#image"
    item = save_evidence_bytes(
        data,
        evidence_dir,
        record["source_id"],
        "source-image",
        1,
        path.suffix,
        source_ref,
        width=width,
        height=height,
        image_format=image_format,
    )
    record["evidence"].append(item)
    record["status"] = "partial"
    record["visual_review_required"] = True
    record["warnings"].append("独立图片必须完成视觉读取后才能标记为已读")


def extract_source(
    path: Path,
    source_id: str,
    evidence_dir: Path,
    pdftoppm: Path | None,
    include_hidden_sheets: bool,
    include_review_sheets: bool,
) -> dict[str, Any]:
    record = base_file_record(source_id, path)
    try:
        suffix = path.suffix.lower()
        if suffix in {".md", ".txt"}:
            extract_text(path, record)
        elif suffix == ".docx":
            extract_docx(path, record, evidence_dir)
        elif suffix == ".pdf":
            extract_pdf(path, record, evidence_dir, pdftoppm)
        elif suffix == ".xlsx":
            extract_xlsx(path, record, evidence_dir, include_hidden_sheets, include_review_sheets)
        elif suffix in IMAGE_EXTENSIONS:
            extract_image(path, record, evidence_dir)
        else:
            record["status"] = "unreadable"
            record["warnings"].append(f"不支持的文件类型：{suffix or '<none>'}")
    except Exception as exc:
        record["status"] = "unreadable"
        record["warnings"].append(f"读取失败：{type(exc).__name__}: {exc}")
    return record


def validate_sources(values: list[str]) -> list[Path]:
    if not values:
        raise SourcePacketError("至少需要一个 --source")
    paths: list[Path] = []
    identities: set[str] = set()
    for value in values:
        path = Path(value).expanduser().resolve()
        if not path.exists() or not path.is_file():
            raise SourcePacketError(f"源文件不存在或不是文件：{path}")
        if path.name.startswith("~$"):
            raise SourcePacketError(f"临时锁文件不能作为源文件：{path.name}")
        identity = os.path.normcase(str(path))
        if identity in identities:
            raise SourcePacketError(f"源文件重复：{path}")
        identities.add(identity)
        paths.append(path)
    return paths


def build_packet(args: argparse.Namespace) -> dict[str, Any]:
    sources = validate_sources(args.source)
    output_dir = args.output_dir.resolve()
    evidence_dir = output_dir / "evidence"
    output_dir.mkdir(parents=True, exist_ok=True)
    evidence_dir.mkdir(parents=True, exist_ok=True)
    pdftoppm = args.pdftoppm.resolve() if args.pdftoppm else None

    files = [
        extract_source(
            path,
            f"SRC-{position:03d}",
            evidence_dir,
            pdftoppm,
            args.include_hidden_sheets,
            args.include_review_sheets,
        )
        for position, path in enumerate(sources, 1)
    ]
    package_sha256 = sha256_bytes(
        canonical_json(
            [{"source_id": item["source_id"], "path": item["path"], "sha256": item["sha256"]} for item in files]
        ).encode("utf-8")
    )
    evidence = [item for file_record in files for item in file_record["evidence"]]
    counts = {
        "files": len(files),
        "readable": sum(item["status"] == "readable" for item in files),
        "partial": sum(item["status"] == "partial" for item in files),
        "unreadable": sum(item["status"] == "unreadable" for item in files),
        "content_units": sum(len(item["content_units"]) for item in files),
        "evidence": len(evidence),
        "excluded_items": sum(len(item["excluded_items"]) for item in files),
        "external_links": sum(len(item["external_links"]) for item in files),
    }
    packet = {
        "schema_version": "1.0",
        "created_at": datetime.now(timezone.utc).isoformat(),
        "package_sha256": package_sha256,
        "options": {
            "include_hidden_sheets": args.include_hidden_sheets,
            "include_review_sheets": args.include_review_sheets,
            "external_links_followed": False,
        },
        "files": files,
        "counts": counts,
    }
    ledger = {
        "schema_version": "1.0",
        "package_sha256": package_sha256,
        "evidence": evidence,
    }
    write_json(output_dir / "source_packet.json", packet)
    write_json(output_dir / "source_evidence_ledger.json", ledger)
    return {"status": "ok", "source_packet": str(output_dir / "source_packet.json"), "counts": counts}


def parse_args() -> argparse.Namespace:
    parser = argparse.ArgumentParser(description="构建 qa-case-xlsx 本地策划案源包")
    parser.add_argument("--source", action="append", default=[], help="本地源文件；可重复")
    parser.add_argument("--output-dir", type=Path, required=True, help="audit 输出目录")
    parser.add_argument("--pdftoppm", type=Path, help="bundled pdftoppm 路径")
    parser.add_argument("--include-hidden-sheets", action="store_true")
    parser.add_argument("--include-review-sheets", action="store_true")
    return parser.parse_args()


def main() -> None:
    try:
        result = build_packet(parse_args())
        print(json.dumps(result, ensure_ascii=False, indent=2))
    except SourcePacketError as exc:
        print(json.dumps({"status": "invalid", "error": str(exc)}, ensure_ascii=False, indent=2))
        raise SystemExit(1) from exc


if __name__ == "__main__":
    main()
